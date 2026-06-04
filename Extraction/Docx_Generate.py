import os
import random
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt

# --- SHARED DATA GENERATORS ---
first_names = ["John", "Jane", "Michael", "Emily", "David", "Sarah", "James", "Jessica", "Robert", "Karen"]
last_names = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Miller", "Davis", "Garcia", "Rodriguez", "Wilson"]

def gen_ssn(): return f"{random.randint(100, 999)}-{random.randint(10, 99)}-{random.randint(1000, 9999)}"
def gen_dob(): return (datetime(1970, 1, 1) + timedelta(days=random.randrange(13149))).strftime("%Y-%m-%d")
def gen_phone(): return f"({random.randint(100, 999)}) {random.randint(100, 999)}-{random.randint(1000, 9999)}"
def gen_email(n): return f"{n.lower().replace(' ', '.')}@{random.choice(['example.com', 'testmail.net', 'mockdata.org'])}"

def add_docx_headers(doc, file_id, current_date):
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(2)
    p_meta.add_run(f"APEX GLOBAL SOLUTIONS INC. | AUDIT DIVISION".ljust(45) + f"RUN DATE: {current_date}".rjust(35))
    
    p_div1 = doc.add_paragraph()
    p_div1.paragraph_format.space_after = Pt(2)
    p_div1.add_run("_" * 85)
    
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_after = Pt(2)
    p_head.add_run(f"DATA DUMP REPORT - SOURCE ID: {file_id}\n" + "=" * 85)
    
    p_cols = doc.add_paragraph()
    p_cols.paragraph_format.space_after = Pt(2)
    p_cols.add_run(f"{'FULL NAME'.ljust(22)}{'SSN'.ljust(13)}{'DOB'.ljust(12)}{'PHONE'.ljust(16)}{'EMAIL'}\n" + "-" * 85)

def create_docx_file(filepath, file_id, target_pages):
    doc = Document()
    
    # Configure global monospace font styling for the Document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(7)
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    records_per_page = 42
    total_records = target_pages * records_per_page - random.randint(5, 15)
    
    page_num = 1
    add_docx_headers(doc, file_id, current_date)
    
    lines_on_page = 0
    for i in range(total_records):
        name = f"{random.choice(first_names)} {random.choice(last_names)}"
        ssn, dob, phone, email = gen_ssn(), gen_dob(), gen_phone(), gen_email(name)
        
        if random.random() < 0.15:
            cp = random.choice(['ssn', 'email', 'phone'])
            if cp == 'ssn': ssn = ssn.replace("-", "")
            elif cp == 'email': email = email.replace("@", "_at_")
            elif cp == 'phone': phone = phone[:5]
            
        row_str = f"{name.ljust(22)}{ssn.ljust(13)}{dob.ljust(12)}{phone.ljust(16)}{email}"
        p_row = doc.add_paragraph()
        p_row.paragraph_format.space_after = Pt(0) # Keep line layouts tight
        p_row.add_run(row_str)
        lines_on_page += 1
        
        if lines_on_page >= records_per_page and i < total_records - 1:
            p_foot = doc.add_paragraph()
            p_foot.add_run("\n" + "_" * 85 + f"\nCONFIDENTIAL RECORD USE ONLY".ljust(60) + f"PAGE {page_num}")
            
            doc.add_page_break() # Structural Word Page Separation
            page_num += 1
            add_docx_headers(doc, file_id, current_date)
            lines_on_page = 0
            
    p_foot = doc.add_paragraph()
    p_foot.add_run("\n" + "_" * 85 + f"\nCONFIDENTIAL RECORD USE ONLY".ljust(60) + f"PAGE {page_num}")
    doc.save(filepath)

if __name__ == "__main__":
    out_dir = "DOCX Extraction"
    os.makedirs(out_dir, exist_ok=True)
    page_targets = [3, 5, 2, 8, 4, 11, 3, 6, 12, 5]
    for i in range(1, 11):
        fid = f"REL{str(i).zfill(6)}"
        create_docx_file(os.path.join(out_dir, f"{fid}.docx"), fid, page_targets[i-1])
    print("Completed: 10 files saved to './DOCX Extraction/'")